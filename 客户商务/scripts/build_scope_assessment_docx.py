#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《研发工作量评估说明》docx：章节骨架对齐招采参考，内容来自总册口径 + 排期说明摘录。

依赖：python-docx；加载报价数据时需 openpyxl（与 build_quotation_xlsx 同模块）。
  .venv_quotation/bin/pip install openpyxl python-docx
  .venv_quotation/bin/python 客户商务/scripts/build_scope_assessment_docx.py
"""
from __future__ import annotations

import importlib.util
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "碳足迹系统_工作量范围评估说明.docx"
REF_SCHEDULE = ROOT / "reference_formats" / "碳足迹-排期说明.docx"


def _load_quotation_module():
    p = ROOT / "scripts" / "build_quotation_xlsx.py"
    spec = importlib.util.spec_from_file_location("carbon_quotation", p)
    m = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(m)
    return m


def _module_day_summary(m) -> list[tuple[str, int]]:
    agg: dict[str, int] = defaultdict(int)
    for r in m.DETAIL_ROWS:
        qid = r[0]
        if qid == "Q-X-002":
            continue
        agg[m.QID_MXX[qid]] += m._DAY_ALLOC[qid]
    keys = sorted(agg.keys())
    return [(m.MXX_TITLE[k], agg[k]) for k in keys]


def _copy_table(dst_doc, src_table):
    nrow, ncol = len(src_table.rows), len(src_table.columns)
    nt = dst_doc.add_table(rows=nrow, cols=ncol)
    nt.style = "Table Grid"
    for ri in range(nrow):
        for ci in range(ncol):
            nt.rows[ri].cells[ci].text = src_table.rows[ri].cells[ci].text
    dst_doc.add_paragraph()


def main():
    m = _load_quotation_module()
    d = Document()

    p0 = d.add_paragraph("产品碳足迹数据服务系统")
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1 = d.add_paragraph("研发工作量评估说明")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

    d.add_heading("1 项目概述", level=1)
    d.add_heading("1.1 项目背景与目标", level=2)
    d.add_paragraph(
        "本项目面向钢铁行业产业链，建设产品碳足迹数据服务系统，支撑订单与任务编排、模板化采集与核算、"
        "第三方核查协作、报告与证书全周期管理等能力，形成可在可信数据空间体系内运营的数字化闭环。"
        "权威业务规则与任务/报告状态以仓库文档 docs/01《整体业务与产品设计审查》、docs/04《任务调度与状态机》、"
        "docs/08《PR 任务与报告边界设计方案》为准。"
    )
    d.add_heading("1.2 分期交付原则", level=2)
    d.add_paragraph(
        "第一阶段聚焦自营 MVP 可运营闭环；第二阶段扩展委托业务、行业 SaaS 与连接器能力；"
        "第三阶段视评审推进轻量化与混合流程。应用模块分层（M01–M10）、分期与范围边界详见"
        "《产品建设方案总册》第一至第六节。"
    )

    d.add_heading("2 项目范围与边界分工", level=1)
    d.add_heading("2.1 建设范围（目标态摘要）", level=2)
    d.add_paragraph(
        "在范围主要包括：多角色工作台；任务五阶段编排；模板与结构化采集；协同澄清；LCA 计算衔接；"
        "核查协作；报告与证书全周期；生态主数据；平台治理与安全基座；分阶段连接与互操作。"
        "不在范围或需单独约定示例：L1 商城本体、生产级中心连接器未单列时、OA 电子签章部分能力等，"
        "与总册 1.3 节及报价编制说明一致。"
    )
    d.add_heading("2.2 相关方与分工（摘要）", level=2)
    d.add_paragraph(
        "荣泽科技：应用研发、集成实施与交付物移交；甲方：需求确认、测试环境及接口配合；"
        "第三方表格控件（SpreadJS）可由客户自采或计入固定价清单；外部 LCA/SaaS 方按接口文档承担对侧开发与联调。"
    )

    d.add_heading("3 项目研发工作量评估", level=1)
    d.add_paragraph(
        "研发工作量涵盖需求分析、方案与原型、编码实现、测试验证、部署运维及相关沟通协调等综合成本。"
        "本文与《碳足迹系统_模块化报价单.xlsx》同源数据：人天单价统一为 1500 元，不按岗位角色拆分单价。"
    )
    d.add_heading("3.1 项目功能模块维度工作量评估", level=2)
    summary = _module_day_summary(m)
    tbl = d.add_table(rows=1 + len(summary), cols=4)
    tbl.style = "Table Grid"
    h = tbl.rows[0].cells
    h[0].text = "编号"
    h[1].text = "应用模块"
    h[2].text = "评估工作量（人天）"
    h[3].text = "说明"
    for i, (name, days) in enumerate(summary, start=1):
        row = tbl.rows[i].cells
        row[0].text = str(i)
        row[1].text = name
        row[2].text = str(days)
        row[3].text = "由报价清单按 M01–M10 汇总，不含 SpreadJS 固定价行"
    d.add_paragraph()

    d.add_heading("3.2 项目人力投入与计价口径", level=2)
    td = m.MVP_LABOR_DAYS + m.PHASE2_OPT_LABOR_DAYS
    ly = td * m.LABOR_RATE
    d.add_paragraph(
        f"人天合计 {td} 人天，按 {m.LABOR_RATE} 元/人天计，人工费用 {ly} 元；"
        f"GrapeCity SpreadJS 商业授权一次性 {m.SPREADJS_FIXED_YUAN} 元（固定价单列）；"
        f"报价总计 {m.TOTAL_QUOTE_YUAN} 元。"
    )

    d.add_heading("4 项目计划里程碑", level=1)
    d.add_paragraph(
        "下列内容由《关于钢铁行业可信数据空间-产品碳足迹数据服务系统的排期说明》摘录，"
        "含附件任务分解表，供与对外排期与任务清单对齐；若与 docs 权威条文冲突，以设计文档为准。"
    )

    if not REF_SCHEDULE.is_file():
        d.add_paragraph(f"（未找到排期说明文件：{REF_SCHEDULE}，请复制到 reference_formats 后重跑。）")
    else:
        sched = Document(str(REF_SCHEDULE))
        for i, para in enumerate(sched.paragraphs):
            if i == 0:
                continue
            t = (para.text or "").strip()
            if not t:
                continue
            st = para.style.name if para.style else "Normal"
            if st == "Title":
                pr = d.add_paragraph(t)
                pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif st.startswith("Heading"):
                lvl = 2
                if "Heading 1" in st:
                    lvl = 2
                elif "Heading 2" in st:
                    lvl = 3
                elif "Heading 3" in st:
                    lvl = 4
                d.add_heading(t, level=min(lvl, 4))
            else:
                d.add_paragraph(t)
        for tbl in sched.tables:
            _copy_table(d, tbl)

    d.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
